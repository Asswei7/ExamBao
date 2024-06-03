from docx import Document
import json


# 1. 创建最后需要上传的word文档
doc = Document()
# 2. 输入当天错题保存的文件数目
for num in range(2):
    with open('0517-{}.json'.format(num+1), 'r', encoding='utf-8') as file:
        # 使用json.load()方法解析JSON数据
        data = json.load(file)
        data = data['data']
    for i in range(len(data)):
        des = str(i + num*50 + 1) + '.' + data[i]['question']
        print(des)
        p = doc.add_paragraph(des+'\n')
        s = data[i]['options']
        jsonS = json.loads(s)
        for item in jsonS:
            s = item['Key'] + '.' + item['Value']
            print(s)
            p = doc.add_paragraph(s+'\n')
        print("答案：" + data[i]['answer'])
        p = doc.add_paragraph("答案：" + data[i]['answer'])


doc.save('17号错题.docx')




